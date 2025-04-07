from rest_framework import viewsets
from .models import Item, Client, Payment, Invoice
from .serializer import ItemSerializer, ClientSerializer, PaymentSerializer, InvoiceSerializer
from django.http import HttpResponse
from django.conf import settings

from docx import Document
import datetime
import pdfkit
import docx
import os

class ItemViewSet(viewsets.ModelViewSet):
    queryset = Item.objects.all()
    serializer_class = ItemSerializer

class ClientViewSet(viewsets.ModelViewSet):
    queryset = Client.objects.all()
    serializer_class = ClientSerializer

class PaymentViewSet(viewsets.ModelViewSet):
    queryset = Payment.objects.all()
    serializer_class = PaymentSerializer

class InvoiceViewSet(viewsets.ModelViewSet):
    queryset = Invoice.objects.all()
    serializer_class = InvoiceSerializer

def generate_invoice(request, invoice_id):
    try:
        # Retrieve the invoice by ID
        invoice = Invoice.objects.get(id=invoice_id)
        client = invoice.client
        items = invoice.items.all()

        # Load the template
        template_path = os.path.join(settings.MEDIA_ROOT, "template.docx")
        document = docx.Document(template_path)

        # # Replace placeholders in the template
        for paragraph in document.paragraphs:
            if "[NAME]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[NAME]", client.supplier_name)
            if "[ADDRESS]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[ADDRESS]", client.address)
            if "[PHONE]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[PHONE]", client.phone_number)
            if "[IN_ID]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[IN_ID]", f"INV-{invoice.id}")
            if "[IN_DATE]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[IN_DATE]", invoice.invoice_date.strftime('%Y-%m-%d'))
            if "[DUE_DATE]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[DUE_DATE]", invoice.due_date.strftime('%Y-%m-%d'))
            if "[STATUS]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[STATUS]", invoice.status)

        # # Add a table for the items
        table = document.tables[0]  # Assuming the first table in the template is for items
        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.item_name
            row_cells[1].text = f"${item.price:.2f}"
            row_cells[2].text = str(item.quantity)

        # # Calculate total amount and add to the document
        total_amount = sum(item.price * item.quantity for item in items)
        total_row = table.add_row().cells
        total_row[0].text = "Total"
        total_row[1].text = f"${total_amount:.2f}"
        total_row[2].text = ""  # Leave quantity cell empty for total row

        # Save the document
        save_path = os.path.join(settings.MEDIA_ROOT, f'invoice_{invoice.id}_{str(datetime.datetime.now().strftime("%d-%m-%Y %HHrs'%M'%S"))}.docx')
        document.save(save_path)

        # Return the file to the client
        with open(save_path, 'rb') as docx_file:
            response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{os.path.basename(save_path)}"'
            return response

    except Invoice.DoesNotExist:
        return HttpResponse("Invoice not found.", status=404)
    except Exception as e:
        return HttpResponse(f"An error occurred: {str(e)}", status=500)